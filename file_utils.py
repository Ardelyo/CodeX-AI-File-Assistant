import os
import shutil
import docx
import re
from rich.progress import Progress

# --- Core File Reading Functions ---
def read_text_file(filepath: str) -> (str | None):
    try:
        if not (os.path.exists(filepath) and os.path.isfile(filepath)): return None
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception: return None

def extract_text_from_pdf(filepath: str) -> (str | None): # Placeholder
    return "PDF parsing not yet fully implemented. Only the filepath was acknowledged."

def extract_text_from_docx(filepath: str) -> (str | None):
    try:
        if not (os.path.exists(filepath) and os.path.isfile(filepath)): return None
        doc = docx.Document(filepath)
        return '\n'.join([para.text for para in doc.paragraphs])
    except Exception: return None

# --- Content Retrieval for General Use (e.g., summarize, ask) ---
def get_file_content(filepath: str, console=None) -> (str | None):
    if not filepath:
        if console: console.print(f"[red]Error: No filepath provided.[/red]")
        return None
    if not os.path.exists(filepath):
        if console: console.print(f"[red]Error: Path not found: '{filepath}'.[/red]")
        return None
    if not os.path.isfile(filepath):
        if console: console.print(f"[red]Error: '{filepath}' is a directory, not a file.[/red]")
        return None

    _, extension = os.path.splitext(filepath.lower())
    content = None
    try:
        if extension in ['.txt', '.py', '.js', '.css', '.html', '.md', '.json', '.xml', '.log', '.ini', '.cfg', '.sh', '.bat']:
            content = read_text_file(filepath)
        elif extension == '.docx':
            content = extract_text_from_docx(filepath)
        elif extension == '.pdf':
            content = extract_text_from_pdf(filepath)
        else:
            if console: console.print(f"[yellow]Unsupported file type for content extraction: {extension} for '{filepath}'[/yellow]")
            return None
        
        if content is None and extension not in ['.pdf']:
             if console: console.print(f"[yellow]Warning: Could not read content from '{filepath}' (empty or error).[/yellow]")
    except Exception as e:
        if console: console.print(f"[red]Error getting content from '{filepath}': {e}[/red]")
    return content

# --- Content Retrieval Optimized for Searching ---
def get_file_content_for_search(filepath: str, console) -> (str | None):
    if not (filepath and os.path.exists(filepath) and os.path.isfile(filepath)): return None
    _, extension = os.path.splitext(filepath.lower())
    content = None
    MAX_SEARCH_CONTENT_SIZE = 100 * 1024  # 100KB
    try:
        if extension in ['.txt', '.py', '.js', '.css', '.html', '.md', '.json', '.xml', '.log', '.ini', '.cfg', '.sh', '.bat']:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read(MAX_SEARCH_CONTENT_SIZE)
        elif extension == '.docx':
            doc = docx.Document(filepath)
            full_text_list = [para.text for para in doc.paragraphs]
            temp_content = ""
            for para_text in full_text_list:
                if len(temp_content) + len(para_text) + 1 > MAX_SEARCH_CONTENT_SIZE:
                    needed = MAX_SEARCH_CONTENT_SIZE - len(temp_content) -1
                    if needed > 0 : temp_content += "\n" + para_text[:needed]
                    break
                temp_content += "\n" + para_text
            content = temp_content.lstrip("\n")
        # elif extension == '.pdf': # Actual PDF parsing for search would go here
            # pass
    except Exception as e:
        if console: console.print(f"[yellow]Warning (search read): Could not get content from '{os.path.basename(filepath)}': {str(e)[:50]}...[/yellow]")
    return content

# --- File and Folder Operations ---
def list_folder_contents(folder_path: str, console) -> (list[dict] | None):
    if not folder_path:
        if console: console.print("[red]Error: No folder path provided for listing.[/red]")
        return None
    abs_folder_path = os.path.abspath(folder_path) # Normalize
    if not os.path.exists(abs_folder_path):
        if console: console.print(f"[red]Error: Folder path '{abs_folder_path}' does not exist.[/red]")
        return None
    if not os.path.isdir(abs_folder_path):
        if console: console.print(f"[red]Error: Path '{abs_folder_path}' is not a directory.[/red]")
        return None
    
    items = []
    try:
        for item_name in os.listdir(abs_folder_path):
            item_path = os.path.join(abs_folder_path, item_name)
            item_type = "file" if os.path.isfile(item_path) else "folder" if os.path.isdir(item_path) else "other"
            items.append({"name": item_name, "type": item_type, "path": item_path})
        return items
    except Exception as e:
        if console: console.print(f"[red]Error listing contents of folder '{abs_folder_path}': {e}[/red]")
    return None

def move_item(source_path: str, destination_path: str, console=None) -> bool:
    abs_source_path = os.path.abspath(source_path)
    abs_destination_path = os.path.abspath(destination_path)

    if not os.path.exists(abs_source_path):
        if console: console.print(f"[red]Error: Source path '{abs_source_path}' does not exist.[/red]")
        return False
    
    try:
        # If destination is an existing directory, move source into it
        if os.path.isdir(abs_destination_path):
            final_dest_path = os.path.join(abs_destination_path, os.path.basename(abs_source_path))
            if os.path.exists(final_dest_path) and final_dest_path != abs_source_path: # Check for overwrite, except if renaming case in same dir
                if console: console.print(f"[red]Error: Item '{os.path.basename(abs_source_path)}' already exists in '{abs_destination_path}'. Move cancelled to prevent overwrite.[/red]")
                return False
        else: # Destination is a full path (could be new name, or new subdirs)
            dest_parent_dir = os.path.dirname(abs_destination_path)
            if not os.path.exists(dest_parent_dir):
                try:
                    os.makedirs(dest_parent_dir, exist_ok=True)
                except OSError as e:
                    if console: console.print(f"[red]Error: Could not create destination directory '{dest_parent_dir}': {e}[/red]")
                    return False
            # Check if trying to move a directory onto an existing file path
            if os.path.isfile(abs_destination_path) and os.path.isdir(abs_source_path):
                if console: console.print(f"[red]Error: Cannot move directory '{abs_source_path}' to replace an existing file '{abs_destination_path}'.[/red]")
                return False
        
        shutil.move(abs_source_path, abs_destination_path)
        return True
    except Exception as e:
        if console: console.print(f"[red]Error moving '{abs_source_path}' to '{abs_destination_path}': {e}[/red]")
    return False

def search_files_recursive(start_path: str, criteria: str, llm_connector, console) -> list:
    found_items = []
    criteria_lower = criteria.lower()
    abs_start_path = os.path.abspath(start_path)

    search_type_keywords = {
        "image": ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp', '.svg', '.heic', '.avif'],
        "picture": ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp', '.svg', '.heic', '.avif'],
        "photo": ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp', '.svg', '.heic', '.avif'],
        "video": ['.mp4', '.mov', '.avi', '.mkv', '.wmv', '.flv', '.webm'],
        "audio": ['.mp3', '.wav', '.ogg', '.aac', '.flac', '.m4a'],
        "document": ['.pdf', '.doc', '.docx', '.odt', '.txt', '.rtf', '.ppt', '.pptx', '.xls', '.xlsx', '.csv', '.md', '.tex'],
        "pdf": [".pdf"], "word document": [".doc", ".docx"], "text file": [".txt", ".md", ".log", ".rtf"],
        "spreadsheet": [".xls", ".xlsx", ".ods", ".csv"], "presentation": [".ppt", ".pptx", ".odp"],
        "python script": [".py", ".pyw"], "javascript file": [".js", ".mjs"], "typescript file": [".ts", ".tsx"],
        "html file": [".html", ".htm"], "css file": [".css", ".scss", ".less"],
        "archive": [".zip", ".rar", ".tar", ".gz", ".7z", ".bz2"], "executable": [".exe", ".msi", ".dmg", ".app", ".deb", ".rpm"],
        "code file": [".py",".js",".java",".c",".cpp",".cs",".go",".rs",".swift",".kt",".php",".rb",".pl",".sh",".bat"],
    }
    
    target_extensions = None
    content_search_term = None
    llm_content_check_criteria = None

    containing_match = re.search(r"containing\s+['\"](.+?)['\"]", criteria_lower)
    if containing_match:
        content_search_term = containing_match.group(1)
        type_description = criteria_lower.split("containing", 1)[0].strip()
    else:
        type_description = criteria_lower

    about_match = re.search(r"(?:about|related to|regarding|on the topic of)\s+['\"](.+?)['\"]", type_description)
    if about_match:
        llm_content_check_criteria = criteria # Pass original full criteria
        type_description = type_description.split(about_match.group(0),1)[0].strip()
    
    if type_description and type_description not in ["files", "any files", "all files", "items"]:
        for key, exts in search_type_keywords.items():
            if key in type_description:
                target_extensions = exts
                break
    
    if not os.path.isdir(abs_start_path):
        console.print(f"[red]Error: Search path '{abs_start_path}' is not a valid directory.[/red]")
        return []
    
    with Progress(console=console, transient=True) as progress_bar:
        search_task_id = progress_bar.add_task("[cyan]Scanning...", total=None)
        for root, dirs, files in os.walk(abs_start_path, topdown=True):
            dirs[:] = [d for d in dirs if not d.startswith('.') and not d.startswith('$')] # Skip hidden/system dirs
            files = [f for f in files if not f.startswith('.')]
            if progress_bar.finished: break
            progress_bar.update(search_task_id, description=f"[cyan]Scanning: {os.path.basename(root)}")
            
            for filename in files:
                if progress_bar.finished: break
                filepath = os.path.join(root, filename)
                _, ext = os.path.splitext(filename.lower())
                if target_extensions and ext not in target_extensions: continue

                item_info = {"name": filename, "type": "file", "path": filepath}
                needs_content_check = content_search_term or llm_content_check_criteria
                
                if needs_content_check:
                    content = get_file_content_for_search(filepath, console)
                    if content:
                        matched_by_simple_search = False
                        if content_search_term and content_search_term in content.lower():
                            found_items.append(item_info)
                            matched_by_simple_search = True
                        
                        if llm_content_check_criteria and not matched_by_simple_search: # Only LLM check if not found by simple search
                            progress_bar.update(search_task_id, description=f"[yellow]LLM check: {filename[:30]}...")
                            if llm_connector.check_content_match(content, llm_content_check_criteria):
                                found_items.append(item_info)
                            progress_bar.update(search_task_id, description=f"[cyan]Scanning: {os.path.basename(root)}")
                else:
                    found_items.append(item_info)
            progress_bar.update(search_task_id, advance=1) # Advance per directory processed
        progress_bar.update(search_task_id, completed=True, description="[green]Search scan complete.")
    return found_items

# __main__ block for testing file_utils.py directly (optional)
if __name__ == '__main__':
    from rich.console import Console as RichConsole # Alias to avoid conflict
    test_console_main = RichConsole()
    test_console_main.print("[bold blue]--- Testing file_utils.py Directly ---[/bold blue]")
    # ... (add specific test cases here if running file_utils.py solo)
    # Example:
    # test_dir_main = "file_utils_direct_test"
    # if not os.path.exists(test_dir_main): os.makedirs(test_dir_main)
    # test_console_main.print(f"Test folder created at: {os.path.abspath(test_dir_main)}")
    # with open(os.path.join(test_dir_main, "test.txt"), "w") as f: f.write("hello world")
    # listed_items = list_folder_contents(test_dir_main, test_console_main)
    # test_console_main.print("Listed items:", listed_items)
    # if os.path.exists(test_dir_main): shutil.rmtree(test_dir_main)
    test_console_main.print("[bold blue]--- End of file_utils.py direct tests ---[/bold blue]")